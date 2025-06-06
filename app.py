from flask import Flask, request, send_file, jsonify
import qrcode
from PIL import Image
import json
import io
import os
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

@app.route("/generate_labels", methods=["POST"])
def generate_labels():
    try:
        data_list = request.get_json().get("data", [])

        # Load doll logo
        logo = Image.open("doll.png")
        logo_size = 100
        logo.thumbnail((logo_size, logo_size))

        # Create Word doc
        doc = Document()
        doc.add_heading("QR Code Labels with Logo", 0)

        for i, entry in enumerate(data_list[:50]):
            qr_content = json.dumps({
                "X1": entry.get("X1", ""),
                "X2": entry.get("X2", ""),
                "X3": entry.get("X3", ""),
                "X4": entry.get("X4", []),
                "X5": entry.get("X5", "")
            }, separators=(',', ':'))

            qr = qrcode.QRCode(
                error_correction=qrcode.constants.ERROR_CORRECT_H,
                box_size=10,
                border=4
            )
            qr.add_data(qr_content)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white").convert("RGB")

            # Insert logo in center
            qr_width, qr_height = img.size
            pos = ((qr_width - logo_size) // 2, (qr_height - logo_size) // 2)
            img.paste(logo, pos, mask=logo if logo.mode == "RGBA" else None)

            # Save temp QR
            temp_qr = f"qr_{i}.png"
            img.save(temp_qr)

            doc.add_paragraph(entry.get("X2", "Label"))
            doc.add_picture(temp_qr, width=Inches(1.5))
            os.remove(temp_qr)

        # Return DOCX
        output = "QR_Code_Labels_with_Logo.docx"
        doc.save(output)
        return send_file(output, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
